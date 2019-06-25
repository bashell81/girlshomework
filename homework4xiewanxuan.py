# -*- coding: UTF-8 -*-
import random
from docx import Document
from docx.shared import Pt

#输出正常乘法1～9
def getChengFa_Normal():
    a = random.choice([1,2,3,4,5,6,7,8,9])
    b = random.choice([1,2,3,4,5,6,7,8,9])
    output = str(a) + " x " + str(b) +" = \t\t"
    return output

#输出第一个填空乘法
def getChengFa_First():
    a = random.choice([1,2,3,4,5,6,7,8,9])
    b = random.choice([1,2,3,4,5,6,7,8,9])
    output = '(   ) x ' + str(b) + ' = ' + str(a*b) + '\t\t'
    return output

def getChengFa_Last():
    a = random.choice([1,2,3,4,5,6,7,8,9])
    b = random.choice([1,2,3,4,5,6,7,8,9])
    output = str(a) + ' x (   ) = ' + str(a*b) + '\t\t'
    return output

def getHomeWorkStr():
    str_ok = ''
    for i in range(80):
        str_ok +=getChengFa_Normal()
        if i%4 == 3:
            str_ok +='\n'

    for i in range(10):
        str_ok +=getChengFa_First()
        if i%4 == 3:
            str_ok +='\n'

    for i in range(2):
        str_ok +=getChengFa_Last()

    str_ok +='\n'
    for i in range(8):
        str_ok +=getChengFa_Last()
        if i%4 == 3:
            str_ok +='\n'
    return str_ok.rstrip('\n')


def write2File(str_oks):
    document = Document()

    for str_ok in str_oks:
        paragraph = document.add_paragraph('')

        # 增加文字
        paragraph.add_run(str_ok)

        paragraph.paragraph_format.line_spacing = Pt(25)

        style = document.styles['Normal']
        font = style.font
        font.size = Pt(14)
        #document.add_page_break()

    document.save('chengfa.docx')

out = []
for x in range(5):
    out.append(getHomeWorkStr())

write2File(out)
